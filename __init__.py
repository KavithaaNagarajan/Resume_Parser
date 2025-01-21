
# scripts/__init__.py

# # Importing functions from various scripts so they are available when the package is imported
# from .Conversion import *                 # Import all functions from Conversion.py
# from .CleanNameFile import *              # Import all functions from CleanNameFile.py
# from .Headername_extraction import *      # Import all functions from Headername_extraction.py
# from .Process_cleanheader import *        # Import all functions from Process_cleanheader.py
# from .Compared_name import *              # Import all functions from Compared_name.py
# from .Email import *                      # Import all functions from Email.py
# from .Phone_number import *               # Import all functions from Phone_number.py
# from .combine_outputs import *            # Import the combine function from combine_outputs.py

# # Optional: Initialize logging setup, versioning, etc.
# import logging

# # Set up basic logging for the package
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Versioning (optional)
# __version__ = '1.0.0'

# # Optional: Any other initialization code (e.g., data loading, settings, etc.)


import os
from PIL import Image
import PyPDF2
import psutil
from spire.doc import Document
from spire.doc import FileFormat

class FileConverter:
    def __init__(self, input_folder, output_folder):
        """
        Initializes the FileConverter class with input and output folder paths.
        
        Args:
        - input_folder (str): The folder containing the files to process.
        - output_folder (str): The folder to save the processed PDF files.
        """
        self.input_folder = input_folder
        self.output_folder = output_folder

        # Make sure output folder exists
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

    def kill_word_processes(self):
        """
        Terminates any running Word processes (winword.exe).
        """
        for proc in psutil.process_iter(['pid', 'name']):
            if proc.info['name'].lower() == 'winword.exe':
                os.kill(proc.info['pid'], 9)
                print(f"Terminated Word process with PID: {proc.info['pid']}")

    def doc_to_pdf(self, doc_path, output_pdf_path):
        """
        Converts a DOC file to PDF using Spire.Doc library.
        """
        try:
            document = Document()
            document.LoadFromFile(doc_path)
            document.SaveToFile(output_pdf_path, FileFormat.PDF)
            document.Close()
            print(f"Successfully converted {doc_path} to {output_pdf_path}")
        except Exception as e:
            print(f"Error converting DOC to PDF: {e}")

    def docx_to_pdf(self, docx_path, output_pdf_path):
        """
        Converts a DOCX file to PDF using Spire.Doc library.
        """
        try:
            document = Document()
            document.LoadFromFile(docx_path)
            document.SaveToFile(output_pdf_path, FileFormat.PDF)
            document.Close()
            print(f"Successfully converted {docx_path} to {output_pdf_path}")
        except Exception as e:
            print(f"Error converting DOCX to PDF: {e}")

    def pdf_to_pdf(self, input_pdf_path, output_pdf_path):
        """
        Copies an existing PDF to ensure consistency or apply other manipulations.
        """
        try:
            with open(input_pdf_path, "rb") as infile:
                reader = PyPDF2.PdfReader(infile)
                writer = PyPDF2.PdfWriter()
                for page_num in range(len(reader.pages)):
                    writer.add_page(reader.pages[page_num])
                with open(output_pdf_path, "wb") as outfile:
                    writer.write(outfile)
            print(f"Copied {input_pdf_path} to {output_pdf_path} successfully.")
        except Exception as e:
            print(f"Error handling PDF: {e}")

    def image_to_pdf(self, image_path, output_pdf_path):
        """
        Converts image files (JPG, PNG, etc.) to PDF.
        """
        try:
            img = Image.open(image_path)
            img.convert("RGB").save(output_pdf_path, "PDF")
            print(f"Converted {image_path} to PDF successfully.")
        except Exception as e:
            print(f"Error converting image to PDF: {e}")

    def convert_to_pdf(self, input_file, output_pdf_path):
        """
        Main function to determine file type and call appropriate conversion function.
        """
        file_extension = input_file.lower().split('.')[-1]
        
        if file_extension in ['doc', 'docx']:
            if file_extension == 'docx':
                self.docx_to_pdf(input_file, output_pdf_path)
            elif file_extension == 'doc':
                self.doc_to_pdf(input_file, output_pdf_path)
        elif file_extension == 'pdf':
            self.pdf_to_pdf(input_file, output_pdf_path)
        elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
            self.image_to_pdf(input_file, output_pdf_path)
        else:
            print(f"Unsupported file format: {file_extension}")

    def process_folder(self):
        """
        Processes all files in the input folder, converts them to PDF, 
        and saves them in the output folder.
        """
        for file_name in os.listdir(self.input_folder):
            file_path = os.path.join(self.input_folder, file_name)
            
            # Skip directories
            if os.path.isdir(file_path):
                continue
            
            # Get the output file path
            output_pdf_path = os.path.join(self.output_folder, f"{os.path.splitext(file_name)[0]}.pdf")
            
            # Convert file to PDF
            print(f"Processing: {file_name}")
            self.convert_to_pdf(file_path, output_pdf_path)
            print(f"Converted {file_name} to PDF and saved as {output_pdf_path}\n")

# Example usage
input_folder = "Data_pass/"  # Replace with the path of your input folder
output_folder = "processed_pdfs/"  # Replace with the desired output folder for PDFs

# Instantiate the FileConverter class
file_converter = FileConverter(input_folder, output_folder)

# Optionally, you can call the method to process the folder
file_converter.process_folder()


import re
import os
import fitz  # PyMuPDF for extracting text from PDFs
import pandas as pd

class FileCleaner:
    def __init__(self, folder_path, csv_file, output_folder):
        """
        Initializes the FileCleaner class with folder paths and CSV file.
        
        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - csv_file (str): Path to the CSV file containing unwanted words.
        - output_folder (str): Path to the folder where cleaned results will be saved.
        """
        self.folder_path = folder_path
        self.csv_file = csv_file
        self.output_folder = output_folder

        # Load unwanted words from the CSV file
        self.unwanted_words = self.load_unwanted_words(self.csv_file)

        # Initialize an empty list to store processed data
        self.data = []

    def load_unwanted_words(self, csv_file):
        """
        Loads unwanted words from a CSV file into a list.
        
        Args:
        - csv_file (str): Path to the CSV file containing unwanted words.
        
        Returns:
        - list: List of unwanted words.
        """
        df = pd.read_csv(csv_file)
        unwanted_words = df['Unwanted Word'].tolist()  # Extract unwanted words from the CSV
        return unwanted_words

    def clean_name_from_filename(self, filename):
        """
        Cleans the file name by removing unwanted words and additional unwanted patterns.
        
        Args:
        - filename (str): The original filename to clean.
        
        Returns:
        - str: Cleaned filename without unwanted words, numbers, special characters, etc.
        """
        name_with_extension = os.path.splitext(filename)[0]
        
        # Create the regex pattern from the unwanted words
        unwanted_pattern = r"(" + "|".join(self.unwanted_words) + r")"  # Join words using OR operator
        
        # Additional cleaning patterns
        additional_patterns = [
            r"\d+",  # Remove numbers
            r"[^\w\s]",  # Remove special characters (brackets, punctuation)
            r"ym$",  # Remove 'ym' from the end of the string
        ]
        
        # Apply the cleaning patterns (unwanted words + additional patterns)
        name_with_extension = re.sub(unwanted_pattern, '', name_with_extension, flags=re.IGNORECASE)
        for pattern in additional_patterns:
            name_with_extension = re.sub(pattern, '', name_with_extension)
        
        # Clean extra spaces
        cleaned_name = re.sub(r"\s+", " ", name_with_extension).strip()
        
        return cleaned_name

    def extract_email(self, text):
        """
        Extract email from text using regex.
        
        Args:
        - text (str): The text to extract the email from.
        
        Returns:
        - str: The extracted email, or None if no email is found.
        """
        email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_regex, text)
        if match:
            return match.group(0)
        return None

    def extract_phone_number_from_text(self, text):
        """
        Extract phone number from text using regex.
        
        Args:
        - text (str): The text to extract phone numbers from.
        
        Returns:
        - str: The extracted phone number, or None if no phone number is found.
        """
        phone_pattern = r'(\+?\(?\d{1,4}\)?[\s\-]?)?(\(?\d{1,4}\)?[\s\-]?\d{1,3}[\s\-]?\d{3}[\s\-]?\d{4}|\d{10}|\+?\d{1,3}[\s\-]?\d{1,4}[\s\-]?\d{1,4}[\s\-]?\d{1,4}|\d{10})'
        phone_matches = re.findall(phone_pattern, text)

        for match in phone_matches:
            phone_number = ''.join(match).replace(' ', '').replace('-', '').replace('(', '').replace(')', '')
            if len(phone_number) >= 10:  # Ensure it's a valid phone number
                return phone_number

        return None

    def process_files(self):
        """
        Processes each PDF file in the folder, extracts and cleans relevant data, 
        and stores the results in a list.
        """
        for filename in os.listdir(self.folder_path):
            if filename.endswith('.pdf'):
                file_path = os.path.join(self.folder_path, filename)
                
                # Extract name from the file name
                cleaned_name = self.clean_name_from_filename(filename)
                
                # Open the PDF and extract email and phone number
                with open(file_path, "rb") as file:
                    pdf_reader = fitz.open(file)
                    text = ""
                    for page_num in range(pdf_reader.page_count):
                        page = pdf_reader.load_page(page_num)
                        text += page.get_text("text")
                
                # Extract email and phone number
                email = self.extract_email(text)
                phone_number = self.extract_phone_number_from_text(text)
                
                # Append the results as a dictionary to the data list
                self.data.append({
                    'Filename': filename,
                    'Name from Filename': cleaned_name,
                    'Email': email,
                    'Phone Number': phone_number
                })

    def save_results(self):
        """
        Converts the processed data into a DataFrame and saves it as a CSV file.
        """
        # Convert the list of dictionaries into a DataFrame
        df = pd.DataFrame(self.data)

        # Ensure the directory for the output file exists
        output_directory = os.path.dirname(self.output_folder)
        
        # Check if the directory exists, and create it if not
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
        
        # Save the cleaned results to a CSV file
        df.to_csv(self.output_folder, index=False)

        print("File cleaning and CSV generation completed successfully.")

# Example usage
folder_path = 'processed_pdfs/'  # Folder where PDF files are stored
csv_file = 'scripts/Filename_unwanted.csv'  # CSV file containing unwanted words
output_folder = 'Prased_resumes/test_trainname_demo.csv'  # Path where the output CSV will be saved

# Instantiate the FileCleaner class
file_cleaner = FileCleaner(folder_path, csv_file, output_folder)

# Process the files and save the results
file_cleaner.process_files()
file_cleaner.save_results()



import os
import spacy
from pdfminer.high_level import extract_text
from docx import Document
import re
import pandas as pd

class HeaderExtractor:
    def __init__(self, folder_path, unwanted_words=None, evaluation_warning=None):
        """
        Initializes the HeaderExtractor with the folder path and optional parameters.

        Args:
        - folder_path (str): The path to the folder containing resumes.
        - unwanted_words (list): List of words to be avoided in the first line.
        - evaluation_warning (str): The warning message to avoid in the first line.
        """
        self.folder_path = folder_path
        self.nlp = spacy.load("en_core_web_sm")  # Load spaCy model for NER
        self.unwanted_words = unwanted_words or [
            "CURRICLUM VITAE", "curriculum vitae", "resume", "contact", 
            "personal details", "contact", "Professional Skills", 
            "Name", "SUMMARY", "SKILLS", "EXPERIENCE"
        ]
        self.evaluation_warning = evaluation_warning or "Evaluation Warning: The document was created with Spire.Doc for Python."
        self.data = []

    def extract_text_from_pdf(self, pdf_path):
        """Extracts text from PDF file."""
        return extract_text(pdf_path)

    def extract_text_from_docx(self, docx_path):
        """Extracts text from DOCX file."""
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def clean_and_get_valid_line(self, text):
        """Cleans the text and returns the first valid line."""
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if line and not any(word.lower() in line.lower() for word in self.unwanted_words) and line != self.evaluation_warning:
                return line
        return "No valid line found."

    def extract_name_using_spacy(self, text):
        """Extracts name using spaCy's Named Entity Recognition (NER)."""
        doc = self.nlp(text)
        names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if names:
            return names[0]
        return None

    def extract_name_using_regex(self, first_line):
        """Extracts name using regex from the first valid line."""
        name_pattern = r'\b([A-Z][a-z]+(?: [A-Z]\.)? [A-Z][a-z]+|[A-Z][a-z]+(?: [A-Z][a-z]+)?)\b'
        matches = re.findall(name_pattern, first_line)
        if matches:
            return matches[0]
        return None

    def extract_full_name(self, resume_path):
        """Extracts full name from the resume, either PDF or DOCX."""
        if resume_path.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(resume_path)
        elif resume_path.lower().endswith('.docx'):
            text = self.extract_text_from_docx(resume_path)
        else:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")
        
        valid_line = self.clean_and_get_valid_line(text)
        full_name = self.extract_name_using_spacy(valid_line)
        if not full_name:
            full_name = self.extract_name_using_regex(valid_line)

        return full_name, valid_line

    def process_folder_and_store_in_dataframe(self):
        """Processes all files in the folder and stores results in a DataFrame."""
        for root, dirs, files in os.walk(self.folder_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    file_path = os.path.join(root, file)
                    print(f"Processing file: {file_path}")
                    full_name, extracted_line = self.extract_full_name(file_path)
                    if full_name:
                        print(f"Candidate's Full Name: {full_name}")
                    else:
                        full_name = "Full name not found."
                    
                    self.data.append({"Filename": file, "Full_Name": full_name, "Extracted_Line": extracted_line})

        df = pd.DataFrame(self.data)
        return df

    def convert_to_proper_case(self, text):
        """Converts text to title case."""
        text = str(text)  # Convert to string if it's a number or NaN
        return text.title()

    def save_results_to_csv(self, output_path):
        """Saves the results DataFrame to a CSV file."""
        # Check if the directory exists, if not, create it
        output_dir = os.path.dirname(output_path)
        if not os.path.exists(output_dir):
            print(f"Directory {output_dir} does not exist. Creating...")
            os.makedirs(output_dir)

        df = pd.DataFrame(self.data)
        
        df['Extracted_Line'] = df['Extracted_Line'].apply(self.convert_to_proper_case)
        df.to_csv(output_path, index=False)
        print(f"Results saved to CSV at {output_path}.")

# Example usage
folder_path = "processed_pdfs/"  # Replace with your folder path
output_path = 'scripts/parsed_headername.csv'  # Path to save the CSV

# Instantiate the HeaderExtractor class
header_extractor = HeaderExtractor(folder_path)

# Process the folder and store results
df = header_extractor.process_folder_and_store_in_dataframe()

# Save the results to a CSV file
header_extractor.save_results_to_csv(output_path)


import pandas as pd
import re
import os

class ResumeCleaner:
    def __init__(self, unwanted_words_csv, input_csv, output_csv):
        """
        Initializes the ResumeCleaner with paths for unwanted words CSV, input CSV, and output CSV.

        Args:
        - unwanted_words_csv (str): Path to the CSV file containing unwanted words.
        - input_csv (str): Path to the CSV file containing the parsed resumes.
        - output_csv (str): Path to save the cleaned resumes.
        """
        self.unwanted_words_csv = unwanted_words_csv
        self.input_csv = input_csv
        self.output_csv = output_csv
        self.unwanted_words = self.read_unwanted_words()
        
        # Check if the input CSV is empty or missing
        if os.path.exists(input_csv):
            try:
                self.df = pd.read_csv(input_csv)
                if self.df.empty:
                    print(f"The input CSV file '{input_csv}' is empty.")
                    self.df = pd.DataFrame()  # Create an empty DataFrame to avoid crashes
            except pd.errors.EmptyDataError:
                print(f"The input CSV file '{input_csv}' is empty or could not be read.")
                self.df = pd.DataFrame()  # Create an empty DataFrame
        else:
            print(f"The input CSV file '{input_csv}' does not exist.")
            self.df = pd.DataFrame()  # Create an empty DataFrame

    def read_unwanted_words(self):
        """Reads unwanted words from the CSV file."""
        if os.path.exists(self.unwanted_words_csv):
            try:
                unwanted_df = pd.read_csv(self.unwanted_words_csv)
                if not unwanted_df.empty:
                    unwanted_words = unwanted_df['word'].tolist()  # Convert the column to a list
                    return unwanted_words
                else:
                    print(f"The unwanted words CSV '{self.unwanted_words_csv}' is empty.")
                    return []
            except pd.errors.EmptyDataError:
                print(f"The unwanted words CSV '{self.unwanted_words_csv}' is empty or could not be read.")
                return []
        else:
            print(f"The unwanted words CSV '{self.unwanted_words_csv}' does not exist.")
            return []

    def clean_extracted_line(self, extracted_line):
        """Cleans the extracted line and keeps only the name."""
        # Remove emails using regex
        extracted_line = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', '', extracted_line)

        # Remove phone numbers (various formats)
        phone_pattern = r'(\+?\(?\d{1,4}\)?[\s\-]?)?(\(?\d{1,4}\)?[\s\-]?\d{1,3}[\s\-]?\d{3}[\s\-]?\d{4}|\d{10}|\+?\d{1,3}[\s\-]?\d{1,4}[\s\-]?\d{1,4}[\s\-]?\d{1,4}|\d{10})'
        extracted_line = re.sub(phone_pattern, '', extracted_line)

        # Remove content inside parentheses (including the parentheses)
        extracted_line = re.sub(r'\(.*?\)', '', extracted_line)

        # Remove all unwanted words
        for word in self.unwanted_words:
            # Make sure to remove the word in a case-insensitive way
            extracted_line = re.sub(r'\b' + re.escape(word) + r'\b', '', extracted_line, flags=re.IGNORECASE)

        # Remove everything except alphabetic words (this helps in keeping potential names)
        extracted_line = re.sub(r'[^a-zA-Z\s]', '', extracted_line)

        # Clean extra spaces and strip leading/trailing spaces
        extracted_line = re.sub(r'\s+', ' ', extracted_line).strip()

        return extracted_line

    def clean_resumes(self):
        """Applies the cleaning function to the 'Extracted Line' column."""
        if not self.df.empty and 'Extracted_Line' in self.df.columns:
            self.df['Cleaned_Extracted_Line'] = self.df['Extracted_Line'].apply(lambda x: self.clean_extracted_line(x))
        else:
            print("No valid 'Extracted_Line' column found or DataFrame is empty.")

    def save_cleaned_resumes(self):
        """Saves the cleaned DataFrame to a new CSV file."""
        if not self.df.empty:
            self.df.to_csv(self.output_csv, index=False)
            print(f"Cleaned resumes saved to {self.output_csv}")
        else:
            print("No data to save to CSV.")

# Example usage:
unwanted_words_csv = 'scripts/Header_unwanted_new.csv'  # Path to unwanted words CSV
input_csv = 'scripts/parsed_headername.csv'  # Path to the previously parsed resumes
output_csv = 'Prased_resumes/cleaned_parsed_headername.csv'  # Path to save the cleaned resumes

# Instantiate the ResumeCleaner class
resume_cleaner = ResumeCleaner(unwanted_words_csv, input_csv, output_csv)

# Clean the resumes
resume_cleaner.clean_resumes()

# Save the cleaned resumes to a CSV file
resume_cleaner.save_cleaned_resumes()


import pandas as pd
import os

class NameComparator:
    def __init__(self, filename_csv, headername_csv, output_csv):
        """
        Initializes the NameComparator with paths for filename CSV, headername CSV, and output CSV.

        Args:
        - filename_csv (str): Path to the CSV file containing filename data.
        - headername_csv (str): Path to the CSV file containing cleaned headername data.
        - output_csv (str): Path to save the comparison results.
        """
        self.filename_csv = filename_csv
        self.headername_csv = headername_csv
        self.output_csv = output_csv

        # Check if the files exist
        if not os.path.exists(filename_csv):
            print(f"File '{filename_csv}' does not exist.")
            self.filename_df = pd.DataFrame()  # Initialize as empty DataFrame
        else:
            try:
                self.filename_df = pd.read_csv(filename_csv)
                if self.filename_df.empty:
                    print(f"The file '{filename_csv}' is empty.")
                    self.filename_df = pd.DataFrame()
            except pd.errors.EmptyDataError:
                print(f"Could not read the file '{filename_csv}' because it is empty or invalid.")
                self.filename_df = pd.DataFrame()

        if not os.path.exists(headername_csv):
            print(f"File '{headername_csv}' does not exist.")
            self.headername_df = pd.DataFrame()  # Initialize as empty DataFrame
        else:
            try:
                self.headername_df = pd.read_csv(headername_csv)
                if self.headername_df.empty:
                    print(f"The file '{headername_csv}' is empty.")
                    self.headername_df = pd.DataFrame()
            except pd.errors.EmptyDataError:
                print(f"Could not read the file '{headername_csv}' because it is empty or invalid.")
                self.headername_df = pd.DataFrame()

    def clean_extracted_line(self):
        """
        Cleans the 'Cleaned_Extracted_Line' column by applying the length condition.
        """
        if not self.headername_df.empty:
            self.headername_df['Cleaned_Extracted_Line'] = self.headername_df['Cleaned_Extracted_Line'].apply(
                lambda x: '' if len(str(x)) >= 3 and len(str(x)) < 10 else x
            )

    def merge_dataframes(self):
        """
        Merges the two DataFrames on the 'Filename' column.
        """
        if not self.filename_df.empty and not self.headername_df.empty:
            combined_df = pd.merge(self.filename_df, self.headername_df, on='Filename', how='inner')
            return combined_df[['Filename', 'Name from Filename', 'Cleaned_Extracted_Line', 'Email', 'Phone Number']]
        else:
            print("One or both DataFrames are empty. Cannot merge.")
            return pd.DataFrame()

    def check_match(self, row):
        """
        Checks and handles the conditions for comparing 'Name from Filename' and 'Cleaned_Extracted_Line'.
        """
        name_from_filename = str(row['Name from Filename']).lower() if pd.notna(row['Name from Filename']) else ''
        cleaned_line = str(row['Cleaned_Extracted_Line']).lower() if pd.notna(row['Cleaned_Extracted_Line']) else ''
        
        # If both columns are NaN, return an empty string
        if pd.isna(row['Name from Filename']) and pd.isna(row['Cleaned_Extracted_Line']):
            return ''
        
        # If 'Name from Filename' is NaN, return 'Cleaned Line'
        if pd.isna(row['Name from Filename']):
            return row['Cleaned_Extracted_Line']
        
        # If 'Cleaned Line' is NaN, return 'Name from Filename'
        if pd.isna(row['Cleaned_Extracted_Line']):
            return row['Name from Filename']
        
        # If 'Cleaned Line' is found in 'Name from Filename' (case-insensitive), return 'Cleaned Line'
        if cleaned_line in name_from_filename:
            return row['Cleaned_Extracted_Line']  # Return Cleaned Line if there's a case-insensitive match
        
        # If no match is found, return 'Name from Filename'
        return row['Name from Filename']

    def compare_names(self):
        """
        Applies the matching logic to the DataFrame and adds the 'Result' column.
        """
        # Merge the DataFrames and check for matches
        merged_df = self.merge_dataframes()
        if not merged_df.empty:
            merged_df['Result'] = merged_df.apply(self.check_match, axis=1)
        else:
            print("No data available to compare.")
            merged_df = pd.DataFrame()
        return merged_df

    def save_comparison_results(self, final_df):
        """
        Saves the comparison results to a CSV file.
        """
        if not final_df.empty:
            final_df.to_csv(self.output_csv, index=False)
            print(f"Comparison results saved to {self.output_csv}")
        else:
            print("No comparison results to save.")

# Example usage:
filename_csv = 'Prased_resumes/test_trainname_demo.csv'  # Path to the filename CSV
headername_csv = 'Prased_resumes/cleaned_parsed_headername.csv'  # Path to the headername CSV
output_csv = 'Prased_resumes/Demo_overall_name.csv'  # Path to save the comparison results

# Instantiate the NameComparator class
name_comparator = NameComparator(filename_csv, headername_csv, output_csv)

# Clean the 'Cleaned_Extracted_Line' column
name_comparator.clean_extracted_line()

# Compare the names and generate the final DataFrame
final_df = name_comparator.compare_names()

# Save the comparison results
name_comparator.save_comparison_results(final_df)


import os
import re
import pandas as pd
from PyPDF2 import PdfReader

class EmailExtractor:
    def __init__(self, folder_path, output_csv_path):
        """
        Initializes the EmailExtractor with folder path and output CSV file path.

        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - output_csv_path (str): Path to save the extracted emails in CSV format.
        """
        self.folder_path = folder_path
        self.output_csv_path = output_csv_path

    @staticmethod
    def extract_first_email(text):
        """
        Extracts the first email address from the given text using regex.

        Args:
        - text (str): Text from which to extract the email.

        Returns:
        - str: The first email found, or None if no email is found.
        """
        email_pattern = r"([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)"
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else None  # Return the first email or None if not found

    @staticmethod
    def extract_text_from_pdf(pdf_file_path):
        """
        Extracts text from a PDF file.

        Args:
        - pdf_file_path (str): Path to the PDF file.

        Returns:
        - str: Extracted text from the PDF.
        """
        text = ""
        try:
            # Read the PDF using PyPDF2
            reader = PdfReader(pdf_file_path)
            for page in reader.pages:
                text += page.extract_text()
        except Exception as e:
            print(f"Error extracting text from {pdf_file_path}: {e}")
        return text

    def extract_first_email_from_folder(self):
        """
        Extracts the first email from each PDF file in the specified folder.

        Returns:
        - pd.DataFrame: DataFrame containing filenames and their corresponding first email.
        """
        result = []
        total_files = 0
        processed_files = 0
        
        # Iterate through all PDF files in the folder
        for filename in os.listdir(self.folder_path):
            total_files += 1
            if filename.endswith('.pdf'):
                pdf_file_path = os.path.join(self.folder_path, filename)
                
                # Extract text from the PDF
                text = self.extract_text_from_pdf(pdf_file_path)
                
                # Extract the first email from the text
                first_email = self.extract_first_email(text)
                
                # If an email is found, store it in the result
                if first_email:
                    result.append({'Filename': filename, 'first_email': first_email})
                processed_files += 1

        # Create a DataFrame from the result
        df = pd.DataFrame(result)

        # Print logging for file processing
        print(f"Total files: {total_files}")
        print(f"Processed files with emails: {processed_files}")
        
        # Return DataFrame with the first email from each PDF
        return df

    def save_emails_to_csv(self, emails_df):
        """
        Saves the DataFrame with extracted emails to a CSV file.

        Args:
        - emails_df (pd.DataFrame): DataFrame containing emails to save.
        """
        emails_df.to_csv(self.output_csv_path, index=False)
        print(f"Results saved to: {self.output_csv_path}")


# Example usage
folder_path = 'processed_pdfs/'  # Path to the folder containing PDF files
output_csv_path = os.path.join(os.path.dirname(__file__), '..', 'Prased_resumes', 'email.csv')  # Path to save the CSV

# Instantiate the EmailExtractor class
email_extractor = EmailExtractor(folder_path, output_csv_path)

# Extract the first email from each PDF file in the folder
emails_df = email_extractor.extract_first_email_from_folder()

# Print the DataFrame with extracted emails
print(emails_df)

# Save the result to a CSV file
email_extractor.save_emails_to_csv(emails_df)



import os
import re
import pandas as pd
from PyPDF2 import PdfReader

class PhoneNumberExtractor:
    def __init__(self, folder_path, output_csv_path):
        """
        Initializes the PhoneNumberExtractor with folder path and output CSV file path.

        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - output_csv_path (str): Path to save the extracted phone numbers in CSV format.
        """
        self.folder_path = folder_path
        self.output_csv_path = output_csv_path

    @staticmethod
    def extract_phone_number_from_text(text):
        """
        Extracts phone numbers from a given text using regex.

        Args:
        - text (str): Text from which to extract the phone numbers.

        Returns:
        - list: A list of valid phone numbers found in the text.
        """
        phone_pattern = r'(\+?\(?\d{1,4}\)?[\s\-]?)?(\(?\d{1,4}\)?[\s\-]?\d{1,3}[\s\-]?\d{3}[\s\-]?\d{4}|\d{10}|\+?\d{1,3}[\s\-]?\d{1,4}[\s\-]?\d{1,4}[\s\-]?\d{1,4}|\d{10})'
        phone_matches = re.findall(phone_pattern, text)

        phone_numbers = []
        for match in phone_matches:
            # Clean phone number by removing non-numeric characters
            phone_number = ''.join(match).replace(' ', '').replace('-', '').replace('(', '').replace(')', '')
            if len(phone_number) >= 10:  # Ensure valid phone number (10 or more digits)
                phone_numbers.append(phone_number)

        return phone_numbers if phone_numbers else None

    @staticmethod
    def extract_text_from_pdf(pdf_file_path):
        """
        Extracts text from a PDF file.

        Args:
        - pdf_file_path (str): Path to the PDF file.

        Returns:
        - str: Extracted text from the PDF.
        """
        text = ""
        try:
            reader = PdfReader(pdf_file_path)
            for page in reader.pages:
                text += page.extract_text()
        except Exception as e:
            print(f"Error extracting text from {pdf_file_path}: {e}")
        return text

    def extract_phone_numbers_from_folder(self):
        """
        Extracts phone numbers from all PDFs in the specified folder.

        Returns:
        - pd.DataFrame: A DataFrame with filenames and corresponding phone numbers.
        """
        result = []
        total_files = 0
        processed_files = 0
        
        # Iterate through all PDF files in the folder
        for filename in os.listdir(self.folder_path):
            total_files += 1
            if filename.endswith('.pdf'):
                pdf_file_path = os.path.join(self.folder_path, filename)
                
                # Extract text from the PDF
                text = self.extract_text_from_pdf(pdf_file_path)
                
                # Extract phone numbers from the text
                phone_numbers = self.extract_phone_number_from_text(text)
                
                # If phone numbers are found, store them in the result
                if phone_numbers:
                    first_phone_number = phone_numbers[0] if phone_numbers else None
                    result.append({'Filename': filename, 'phone_numbers': phone_numbers, 'first_phone_number': first_phone_number})
                processed_files += 1

        # Create a DataFrame from the result
        df = pd.DataFrame(result)

        # Print logging for file processing
        print(f"Total files: {total_files}")
        print(f"Processed files with phone numbers: {processed_files}")
        
        # Return the DataFrame
        return df

    def process_phone_numbers(self, row):
        """
        Processes phone numbers based on the provided logic and rules.

        Args:
        - row (pd.Series): A row from the DataFrame containing phone numbers.

        Returns:
        - str: The processed phone number.
        """
        first_phone_number = row['first_phone_number']
        phone_numbers = row['phone_numbers']
        
        # Clean the phone numbers (remove non-numeric characters)
        cleaned_phone_numbers = [''.join(filter(str.isdigit, num)) for num in phone_numbers]
        
        # Apply the logic based on the value of first_phone_number
        if len(first_phone_number) == 4:  # Logic 1
            return cleaned_phone_numbers[-1][-10:] if len(cleaned_phone_numbers[-1]) >= 10 else None
        elif len(first_phone_number) > 13:  # Logic 2
            return cleaned_phone_numbers[-1][-10:] if len(cleaned_phone_numbers[-1]) >= 10 else None
        elif len(first_phone_number) == 6:  # Logic 3
            return cleaned_phone_numbers[-1][-10:] if len(cleaned_phone_numbers[-1]) >= 10 else None
        elif len(first_phone_number) >= 2:  # Logic 4
            return cleaned_phone_numbers[0]  # Take the first phone number
        else:
            return None

    def save_phone_numbers_to_csv(self, phone_numbers_df):
        """
        Saves the DataFrame containing phone numbers to a CSV file.

        Args:
        - phone_numbers_df (pd.DataFrame): DataFrame containing phone numbers to save.
        """
        if not phone_numbers_df.empty:
            phone_numbers_df.to_csv(self.output_csv_path, index=False)
            print(f"Results saved to: {self.output_csv_path}")
        else:
            print("No phone numbers to save.")

# Example usage
folder_path = 'processed_pdfs/'  # Path to the folder containing PDF files
output_csv_path = os.path.join(os.path.dirname(__file__), '..', 'Prased_resumes', 'phone_numbers.csv')  # Path to save the CSV

# Instantiate the PhoneNumberExtractor class
phone_number_extractor = PhoneNumberExtractor(folder_path, output_csv_path)

# Extract phone numbers from PDFs in the folder
phone_numbers_df = phone_number_extractor.extract_phone_numbers_from_folder()

# Check if the DataFrame is not empty before attempting to process phone numbers
if not phone_numbers_df.empty:
    # Process phone numbers
    phone_numbers_df['processed_phone_number'] = phone_numbers_df.apply(phone_number_extractor.process_phone_numbers, axis=1)
    
    # Save the phone numbers to CSV
    phone_number_extractor.save_phone_numbers_to_csv(phone_numbers_df)
else:
    print("No phone numbers found in the folder.")


import os
import pandas as pd

class DataMerger:
    def __init__(self, phone_numbers_csv, email_csv, name_csv, output_csv_path):
        """
        Initializes the DataMerger with paths for the input CSV files and the output CSV path.

        Args:
        - phone_numbers_csv (str): Path to the CSV file containing phone numbers data.
        - email_csv (str): Path to the CSV file containing email data.
        - name_csv (str): Path to the CSV file containing name data.
        - output_csv_path (str): Path to save the merged data as a CSV.
        """
        self.phone_numbers_csv = phone_numbers_csv
        self.email_csv = email_csv
        self.name_csv = name_csv
        self.output_csv_path = output_csv_path

        # Read the CSV files into DataFrames
        self.phone_numbers_df = pd.read_csv(phone_numbers_csv)
        self.email_df = pd.read_csv(email_csv)
        self.name_df = pd.read_csv(name_csv)

    def clean_and_prepare_data(self):
        """
        Cleans and prepares the DataFrames for merging by renaming columns and selecting relevant ones.
        """
        # Clean the 'phone_numbers_df'
        self.phone_numbers_df = self.phone_numbers_df[['Filename', 'phone_numbers', 'first_phone_number']]
        self.phone_numbers_df.rename(columns={'phone_numbers': 'PhoneNumbers_List', 
                                              'first_phone_number': 'Candidates_PhoneNumber'}, inplace=True)

        # Clean the 'email_df'
        self.email_df.rename(columns={'first_email': 'Candidates_Email'}, inplace=True)

        # Clean the 'name_df'
        self.name_df.rename(columns={'Name from Filename': 'Name_from_File',
                                      'Cleaned_Extracted_Line': 'Header_Name', 
                                      'Result': 'Candidates_Name', 
                                      'Email': 'Header_Emails', 
                                      'Phone Number': 'Header_Phone'}, inplace=True)
        self.name_df = self.name_df[['Filename', 'Candidates_Name', 'Name_from_File', 'Header_Name', 'Header_Emails', 'Header_Phone']]

    def merge_dataframes(self):
        """
        Merges the three DataFrames on the 'filename' column using left joins.
        """
        # Perform left joins to merge the three dataframes based on the 'filename' column
        merged_df = self.name_df.merge(self.email_df, on='Filename', how='left') \
                                .merge(self.phone_numbers_df, on='Filename', how='left')

        # Apply the logic for 'Header_Emails' column
        merged_df['Header_Emails'] = merged_df['Header_Emails'].fillna(merged_df['Candidates_Email'])
        
        # Apply the logic for 'Candidates_PhoneNumber' column
        merged_df['Candidates_PhoneNumber'] = merged_df['Candidates_PhoneNumber'].fillna(merged_df['Header_Phone'])

        # Clean up and rename columns
        merged_df.rename(columns={'Header_Emails': 'Candidates_Emails'}, inplace=True)
        merged_df = merged_df[['Filename', 'Candidates_Name', 'Name_from_File', 'Header_Name', 'Candidates_Emails', 'Candidates_PhoneNumber', 'PhoneNumbers_List']]

        return merged_df

    def save_to_csv(self, merged_df):
        """
        Saves the merged DataFrame to a CSV file.

        Args:
        - merged_df (pd.DataFrame): The DataFrame to save.
        """
        merged_df.to_csv(self.output_csv_path, index=False)
        print(f"Combined results saved to: {self.output_csv_path}")

    def process_data(self):
        """
        Orchestrates the cleaning, merging, and saving of the data.
        """
        # Clean and prepare the data
        self.clean_and_prepare_data()

        # Merge the DataFrames
        merged_df = self.merge_dataframes()

        # Save the result to CSV
        self.save_to_csv(merged_df)


# Example usage
phone_numbers_csv = os.path.join('Prased_resumes', 'phone_numbers.csv')
email_csv = os.path.join('Prased_resumes', 'email.csv')
name_csv = os.path.join('Prased_resumes', 'Demo_overall_name.csv')
output_csv_path = os.path.join('Prased_resumes', 'overall_Details.csv')

# Instantiate the DataMerger class
data_merger = DataMerger(phone_numbers_csv, email_csv, name_csv, output_csv_path)

# Process and save the merged data
data_merger.process_data()


import re
import os
import pandas as pd

class ExperienceExtractor:
    def __init__(self, folder_path, output_csv_path):
        """
        Initialize the ExperienceExtractor class.

        Args:
            folder_path (str): Path to the folder containing the PDF files.
            output_csv_path (str): Path to save the output CSV file containing experience data.
        """
        self.folder_path = folder_path
        self.output_csv_path = output_csv_path
        self.experience_data = []

    def extract_experience_from_filename(self, filename):
        """
        Extract total experience from a filename using regex patterns.

        Args:
            filename (str): The name of the file from which to extract experience.

        Returns:
            tuple: A tuple of (years, months) representing total experience.
        """
        total_experience = None
        
        # Pattern 1: Numbers before "year", "years", "yr", "yrs" or "m" (for experience)
        pattern_1 = r'(\d+(\.\d+)?)\s*(year|years|yr|yrs|m)'
        match_1 = re.search(pattern_1, filename, re.IGNORECASE)
        
        if match_1:
            years = None
            months = None
            
            # Check if the match contains floating point years
            years_value = float(match_1.group(1))  # Convert the value to float
            
            # If there's a decimal, take the integer part as years and the decimal part as months
            if years_value.is_integer():
                years = int(years_value)
                months = 0
            else:
                years = int(years_value)  # Take the integer part as years
                # Subtract 1 month from the fractional part and convert to months
                months = int((years_value - years) * 12) - 1  # Subtract 1 month to match your requirement
            
            total_experience = (years, months)

        # Pattern 2: Numbers in square brackets like [10y_8m]
        pattern_2 = r'\[(\d+)(y|yrs|yr?)?[_-](\d+)(m)?\]'
        match_2 = re.search(pattern_2, filename)
        
        if match_2:
            years_in_bracket = int(match_2.group(1))
            months_in_bracket = int(match_2.group(3))
            total_experience = (years_in_bracket, months_in_bracket)

        return total_experience

    def process_files_in_folder(self):
        """
        Iterate through files in the specified folder and extract experience information.

        Returns:
            list: A list of lists containing filenames and their corresponding experience.
        """
        # Iterate over all files in the directory
        for filename in os.listdir(self.folder_path):
            if filename.endswith(".pdf"):  # Process only PDF files
                experience = self.extract_experience_from_filename(filename)
                if experience:  # If experience is found
                    # Append the result to the list as a tuple (filename, experience)
                    self.experience_data.append([filename, f"{experience[0]} years, {experience[1]} months"])
    
    def save_experience_to_csv(self):
        """
        Save the extracted experience data to a CSV file.

        Returns:
            None
        """
        # Convert the result into a pandas DataFrame for easy viewing
        df = pd.DataFrame(self.experience_data, columns=["Filename", "total_exp"])

        # Optionally, save the DataFrame to a CSV file
        df.to_csv(self.output_csv_path, index=False)

        # Print confirmation message
        print(f"Experience data has been saved to '{self.output_csv_path}'")

    def extract_and_save(self):
        """
        Extract experience data from files in the folder and save it to CSV.

        Returns:
            None
        """
        # Process the files and get experience data
        self.process_files_in_folder()

        # Save the data to CSV
        self.save_experience_to_csv()

# Example usage of the ExperienceExtractor class

# Path to the folder containing PDF files
folder_path = 'processed_pdfs/'

# Path to save the output CSV file
output_csv_path = 'Prased_resumes/Fileexp_data.csv'

# Instantiate the class
experience_extractor = ExperienceExtractor(folder_path, output_csv_path)

# Extract experience data and save to CSV
experience_extractor.extract_and_save()



import os
import re
from PyPDF2 import PdfReader
import pandas as pd

class RandomExperienceExtractor:
    def __init__(self, folder_path, output_csv_path):
        # Initialize the class with the folder path and output CSV path
        self.folder_path = folder_path
        self.output_csv_path = output_csv_path
    
    # Method to extract experience from a single PDF
    def extract_experience_from_pdf(self, file_path):
        try:
            # Read the PDF
            reader = PdfReader(file_path)
            if len(reader.pages) == 0:
                return None
            
            # Extract text from the first page
            first_page_text = reader.pages[0].extract_text()

            # Regular expression to capture years of experience including:
            experience_pattern = r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+of(?:\s+[a-zA-Z]+)*\s*(experience|progressively\s+in\s+[a-zA-Z\s]+|professional\s+experience|rich\s+and\s+extensive\s+experience)\b"

            # Matches for experience including keywords like "years", "Yrs", "progressively", etc.
            matches = re.findall(experience_pattern, first_page_text, re.IGNORECASE)

            # Matches for experience in the context of "professional practice"
            additional_experience_pattern = r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+(?:of\s+[a-zA-Z]+)*\s+professional\s+practice\b"
            additional_matches = re.findall(additional_experience_pattern, first_page_text, re.IGNORECASE)

            # Combine both matches
            total_matches = matches + additional_matches

            # Return the unique experience matches as years (remove duplicates)
            return list(set([match[0] for match in total_matches]))  # Extract only the numeric values and remove duplicates
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return None

    # Method to process all PDFs in the folder and extract experience
    def process_folder(self):
        data = []
        for filename in os.listdir(self.folder_path):
            if filename.endswith('.pdf'):
                file_path = os.path.join(self.folder_path, filename)
                experiences = self.extract_experience_from_pdf(file_path)
                
                # Convert the list of experiences into a comma-separated string
                experience_str = ', '.join(experiences) if experiences else "No experience found"
                
                data.append({"Filename": filename, "experience_list": experience_str})
        
        # Create a DataFrame without any sum or rounding, keeping raw experiences as they appear
        df = pd.DataFrame(data)
        return df

    # Method to filter out rows with no experience and apply experience conversion
    def filter_and_convert_experience(self, df):
        # Filter out rows where 'experience_list' is 'No experience found'
        filtered_df = df[df['experience_list'] != 'No experience found']
        
        # Function to convert experience to the required format ("X years, Y months")
        def convert_experience_to_year_month(exp_str):
            try:
                # Handle the case where experience is already a decimal number (e.g., 7.4 or 7.10)
                exp_num = float(exp_str)  # Convert to float for easy manipulation
                years = int(exp_num)  # Integer part is years
                months = round((exp_num - years) * 12)  # Decimal part converted to months
                return f"{years} years, {months} months"
            except ValueError:
                pass  # Continue to the next case if it cannot be converted to float
            
            # Regex patterns to match experience formats like "X years", "X.XX years", "X years Y months"
            year_month_pattern = r'(\d+)(?:\s?years?|yr)?(?:\s?(\d+)\s?months?)?'

            # Try to match the pattern
            match = re.match(year_month_pattern, str(exp_str).lower())

            if match:
                years = int(match.group(1))  # Always has years
                months = int(match.group(2)) if match.group(2) else 0  # Optional months part
                return f"{years} years, {months} months"

            # If nothing matches, assume it's a simple number of years without months
            try:
                years = int(exp_str)
                return f"{years} years, 0 months"
            except ValueError:
                return None  # In case it doesn't match any expected format

        # Apply the conversion function to the 'experience_list' column
        filtered_df['formatted_experience'] = filtered_df['experience_list'].apply(convert_experience_to_year_month)
        
        return filtered_df

    # Method to save the processed experience data to a CSV file
    def save_experience_to_csv(self, filtered_df):
        filtered_df.to_csv(self.output_csv_path, index=False)
        print(f"Experience data saved to {self.output_csv_path}")

    # Method to run the experience extraction and conversion process
    def extract_and_process_experience(self):
        # Step 1: Extract experience data from all PDFs in the folder
        filtered_df = self.process_folder()
        
        # Step 2: Filter and convert experience data
        filtered_df_filtered = self.filter_and_convert_experience(filtered_df)
        
        # Step 3: Save the processed experience data to CSV
        self.save_experience_to_csv(filtered_df_filtered)

# Example usage
folder_path = 'processed_pdfs/'  # Change this to your folder path containing PDF resumes
output_csv_path = 'Prased_resumes/Randomexp_data.csv'  # Path where the output CSV should be saved

# Create an instance of ExperienceExtractor
random_experience_extractor = RandomExperienceExtractor(folder_path, output_csv_path)

# Extract and process experience data
random_experience_extractor.extract_and_process_experience()

import fitz  # PyMuPDF
import pandas as pd
import os
import re

class ResumeExperienceExtractor:
    def __init__(self, folder_path, search_words_file, output_csv_path):
        """
        Initialize the class with the necessary parameters.

        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - search_words_file (str): Path to the CSV file containing search words.
        - output_csv_path (str): Path where the output CSV file will be saved.
        """
        self.folder_path = folder_path
        self.search_words_file = search_words_file
        self.output_csv_path = output_csv_path
        self.search_words = self.load_search_words(search_words_file)
    
    def load_search_words(self, csv_file):
        """
        Load search words from a CSV file and return them as a list.

        Args:
        - csv_file (str): Path to the CSV file containing the search words.

        Returns:
        - list: List of search words extracted from the 'Search_words' column.
        """
        search_words_df = pd.read_csv(csv_file)
        return search_words_df['Search_words'].tolist()

    def extract_experience(self, text):
        """
        Extract years of experience from the provided text based on predefined patterns and custom search words.

        Args:
        - text (str): The extracted text from a PDF.

        Returns:
        - float: Total years of experience found in the text.
        """
        # Define the patterns for experience extraction
        experience_pattern = r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+of(?:\s+[a-zA-Z]+)*\s*(experience|progressively\s+in\s+[a-zA-Z\s]+|professional\s+experience|rich\s+and\s+extensive\s+experience)\b"
        additional_experience_pattern = r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+(?:of\s+[a-zA-Z]+)*\s+professional\s+practice\b"
        
        total_exp = 0

        # Search using the patterns
        patterns = [experience_pattern, additional_experience_pattern]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            
            # If matches are found, process the number of years
            if matches:
                for match in matches:
                    try:
                        exp_years = float(match[0])  # The first group is the number of years
                        total_exp += exp_years
                    except ValueError:
                        continue
        
        # Optionally, search for additional search words in the text
        for word in self.search_words:
            pattern = r'(\d{1,2})\s*(years?|Yrs?)\s*(?=\s+' + re.escape(word) + ')'
            matches = re.findall(pattern, text, re.IGNORECASE)
            
            if matches:
                for match in matches:
                    try:
                        exp_years = int(match[0])  # Extract years from matches
                        total_exp += exp_years
                    except ValueError:
                        continue

        return total_exp

    def process_pdfs(self):
        """
        Process all PDF files in the given folder and extract total years of experience.

        Returns:
        - pd.DataFrame: A DataFrame containing filenames and their corresponding total experience.
        """
        data = []

        # Loop over all the files in the folder
        for filename in os.listdir(self.folder_path):
            if filename.lower().endswith(".pdf"):
                file_path = os.path.join(self.folder_path, filename)
                doc = fitz.open(file_path)
                
                # Extract all text from the PDF
                text = ""
                for page in doc:
                    text += page.get_text("text")
                
                # Extract the experience from the text using the custom patterns
                total_exp = self.extract_experience(text)
                
                # Append to data list (filename, total experience)
                data.append({
                    'Filename': filename,
                    'total_Exp': total_exp
                })
        
        # Convert the data to a DataFrame
        df = pd.DataFrame(data)
        return df

    def save_to_csv(self, df):
        """
        Save the processed DataFrame to a CSV file.

        Args:
        - df (pd.DataFrame): The DataFrame containing experience data.
        """
        df.to_csv(self.output_csv_path, index=False)
        print(f"Results saved to {self.output_csv_path}")

    def run(self):
        """
        Run the entire process: load search words, process PDFs, filter data, and save to CSV.
        """
        # Process PDFs and get the DataFrame
        df = self.process_pdfs()

        # Filter out rows where total_Exp is zero
        df = df[df['total_Exp'] != 0.0]

        # Save the results to the output CSV file
        self.save_to_csv(df)


# Example Usage
folder_path = 'processed_pdfs/'  # Path to your folder containing PDFs
search_words_file = 'scripts/CSV_serachExpheader.csv'  # CSV file containing search words
output_csv_path = 'Prased_resumes/csv_output_experience.csv'  # Path to save the results

# Create an instance of the class
resume_extractor = ResumeExperienceExtractor(folder_path, search_words_file, output_csv_path)

# Run the process
resume_extractor.run()


import pandas as pd
import numpy as np
import re

class ExperienceProcessor:
    def __init__(self, filename_csv, search_csv, content_csv, output_csv):
        """
        Initializes the ExperienceProcessor class with paths to the input and output CSV files.
        
        Args:
        - filename_csv (str): Path to the file containing filename and experience data.
        - search_csv (str): Path to the file containing search words.
        - content_csv (str): Path to the file containing content experience data.
        - output_csv (str): Path to save the final output CSV file.
        """
        self.filename_df = self.read_and_clean_csv(filename_csv)
        self.search_df = self.read_and_clean_csv(search_csv)
        self.content_df = self.read_and_clean_csv(content_csv)
        self.output_csv_path = output_csv

    # Function to read CSV files and clean column names
    def read_and_clean_csv(self, file_path):
        """
        Reads a CSV file and strips any leading/trailing whitespace from column names.
        
        Args:
        - file_path (str): Path to the CSV file.
        
        Returns:
        - pd.DataFrame: DataFrame with cleaned column names.
        """
        df = pd.read_csv(file_path)
        df.columns = df.columns.str.strip()  # Clean column names
        return df

    # Function to merge multiple DataFrames on the 'filename' column
    def merge_dataframes(self):
        """
        Merges the three DataFrames (filename_df, search_df, content_df) on the 'filename' column.
        
        Returns:
        - pd.DataFrame: The merged DataFrame.
        """
        merge_df = pd.merge(self.filename_df, self.search_df, on='Filename', how='outer')
        merged_df = pd.merge(merge_df, self.content_df, on='Filename', how='outer')
        print(merged_df.columns)
        return merged_df

    # Function to convert "years, months" to decimal years (in 1 decimal place)
    def convert_to_decimal_years(self, exp):
        """
        Converts experience in 'years, months' format to decimal years (1 decimal place).
        
        Args:
        - exp (str): The experience string (e.g., '2 years, 6 months').
        
        Returns:
        - float: Experience in decimal years, rounded to 1 decimal place.
        """
        exp = str(exp)
        year_month_pattern = r"(\d+)\s*years?,?\s*(\d+)\s*months?|(\d+)\s*years?"
        match = re.search(year_month_pattern, exp)
        
        if match:
            if match.group(1) and match.group(2):
                years = int(match.group(1))
                months = int(match.group(2))
                decimal_years = years + (months / 12)
            elif match.group(3):
                years = int(match.group(3))
                decimal_years = years
            
            return round(decimal_years, 1)
        
        return None  # Return None if the format doesn't match

    # Function to apply conversion of experience to decimal years
    def apply_decimal_conversion(self, df, column_name):
        """
        Applies the decimal conversion function to a specific column in the DataFrame.
        
        Args:
        - df (pd.DataFrame): DataFrame containing the experience column.
        - column_name (str): Column name containing the experience data.
        
        Returns:
        - pd.DataFrame: DataFrame with the new decimal experience column.
        """
        df[f'{column_name}_decimal'] = df[column_name].apply(self.convert_to_decimal_years)
        return df

    # Function to assign value to 'overall_exp' based on available experience
    def get_overall_exp(self, row):
        """
        Assigns a value to 'overall_exp' based on priority: Filenme_exp, Content_exp, Header_exp.
        
        Args:
        - row (pd.Series): A row from the DataFrame.
        
        Returns:
        - float or np.nan: The assigned value for overall_exp.
        """
        if pd.notna(row['Filenme_exp']):
            return row['Filenme_exp']
        elif pd.notna(row['Content_exp']):
            return row['Content_exp']
        elif pd.notna(row['Header_exp']):
            return row['Header_exp']
        else:
            return np.nan

    # Function to generate the final DataFrame and save it to a CSV file
    def generate_final_dataframe(self):
        """
        Merges input DataFrames, processes experience data, and saves the result to a CSV file.
        """
        # Merge the DataFrames
        merged_df = self.merge_dataframes()
        
        # Apply conversion for 'total_Exp_x' column
        merged_df = self.apply_decimal_conversion(merged_df, 'total_exp')
        
        # Rename columns to match required format
        merged_df = merged_df[['Filename', 'total_exp', 'total_Exp', 'experience_list']]
        merged_df.rename(columns={'Filename': 'Filename', 
                                  'total_exp': 'Filenme_exp', 
                                  'total_Exp': 'Header_exp', 
                                  'experience_list': 'Content_exp'}, inplace=True)
        
        # Apply function to get 'overall_exp'
        merged_df['overall_exp'] = merged_df.apply(self.get_overall_exp, axis=1)
        
        # Save the final DataFrame to a CSV
        merged_df = merged_df[['Filename', 'overall_exp']]
        merged_df.to_csv(self.output_csv_path, index=False)
        print(f"Results saved to {self.output_csv_path}")

# Example Usage
def main():
    # Initialize the ExperienceProcessor class with file paths
    processor = ExperienceProcessor(
        filename_csv='Prased_resumes/experience_data.csv',
        search_csv='Prased_resumes/Randomexp_data.csv',
        content_csv='Prased_resumes/csv_output_experience.csv',
        output_csv='Prased_resumes/Compare_exp.csv'
    )
    
    # Generate the final DataFrame and save it
    processor.generate_final_dataframe()

# Run the main function
if __name__ == "__main__":
    main()


import os
import pandas as pd
import pdfplumber
from fuzzywuzzy import fuzz

class PDFMatcher:
    def __init__(self, pdf_folder_path, jd_csv_path, output_csv_path):
        """
        Initializes the PDFMatcher class with necessary parameters.

        Args:
            pdf_folder_path (str): Path to the folder containing PDF files.
            jd_csv_path (str): Path to the CSV file containing job descriptions and skills.
            output_csv_path (str): Path to save the results as a CSV file.
        """
        self.pdf_folder_path = pdf_folder_path
        self.jd_csv_path = jd_csv_path
        self.output_csv_path = output_csv_path
        self.results = []

    def extract_text_from_pdf(self, pdf_path):
        """
        Extracts text from a PDF file and removes a specific sentence if it exists.

        Args:
            pdf_path (str): Path to the PDF file.

        Returns:
            str: Extracted text from the PDF.
        """
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text()

        # Remove the specific sentence if it exists in the text
        sentence_to_remove = "Evaluation Warning: The document was created with Spire.Doc for Python."
        text = text.replace(sentence_to_remove, "")

        return text

    def calculate_match_percentage(self, text, search_word):
        """
        Calculate the fuzzy match percentage between the extracted text and a search word.

        Args:
            text (str): The text from the PDF.
            search_word (str): The word to match in the text.

        Returns:
            int: The match percentage between the word and the text.
        """
        if not text or not search_word:
            return 0
        # Use fuzzywuzzy to find the match percentage
        match_score = fuzz.partial_ratio(text.lower(), search_word.lower())
        return match_score

    def process_pdfs_and_find_matches(self):
        """
        Processes PDFs in a folder and matches them with skills from a provided job description CSV.

        Returns:
            pandas.DataFrame: DataFrame containing the match results.
        """
        # Load the SQL JD CSV file
        sql_jd_df = pd.read_csv(self.jd_csv_path)

        # Iterate over each PDF file in the folder
        for filename in os.listdir(self.pdf_folder_path):
            if filename.endswith('.pdf'):
                pdf_path = os.path.join(self.pdf_folder_path, filename)
                # Extract text from the current PDF
                pdf_text = self.extract_text_from_pdf(pdf_path)

                # Iterate over each search word in sql_JD file
                for _, row in sql_jd_df.iterrows():
                    search_word = row['Skills']
                    role = row['Roles']  # Assuming the 'Roles' column exists in your CSV

                    # Find the matching percentage for this search word
                    match_percentage = self.calculate_match_percentage(pdf_text, search_word)

                    # If there's a match (match percentage > 0, you can add a threshold like 50%)
                    if match_percentage > 0:
                        self.results.append({
                            'Filename': filename,
                            'matchword': search_word,
                            'matchpercentage': match_percentage,
                            'role': role  # Add the role to the result
                        })

        # Create a DataFrame from the results
        results_df = pd.DataFrame(self.results)

        # Save the results to a CSV if the DataFrame is not empty
        if not results_df.empty:
            results_df.to_csv(self.output_csv_path, index=False)
            print(f"Matches found. Results saved to '{self.output_csv_path}'.")
        else:
            print("No matches found.")

        return results_df

# Example usage
def main():
    # Set the folder path where the resumes (PDFs) are stored
    pdf_folder_path = 'processed_pdfs/'

    # Path to the SQL JD CSV file containing job descriptions and skills
    jd_csv_path = 'scripts/All_JD.csv'

    # Set the output path for the results CSV
    output_csv_path = 'Prased_resumes/Skill_match_results.csv'

    # Create an instance of the PDFMatcher class
    pdf_matcher = PDFMatcher(pdf_folder_path, jd_csv_path, output_csv_path)

    # Process the PDFs and find matches
    results_df = pdf_matcher.process_pdfs_and_find_matches()

    # Optionally, print or manipulate the results DataFrame
    if not results_df.empty:
        print(results_df)

# Run the main function
if __name__ == "__main__":
    main()










